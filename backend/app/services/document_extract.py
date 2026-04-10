from __future__ import annotations

import io
import json
import logging
import re
import xml.etree.ElementTree as ET
import zipfile
from typing import Final

logger = logging.getLogger(__name__)

# Extensions we attempt to parse (lowercase, with dot)
SUPPORTED_EXTENSIONS: Final[frozenset[str]] = frozenset(
    {
        ".pdf",
        ".docx",
        ".doc",
        ".txt",
        ".text",
        ".md",
        ".markdown",
        ".html",
        ".htm",
        ".xml",
        ".json",
        ".rtf",
        ".csv",
        ".tsv",
        ".xls",
        ".xlsx",
        ".pages",
        ".page",
    }
)


def normalize_extension(filename: str) -> str:
    lower = filename.lower().strip()
    if "." not in lower:
        return ""
    return "." + lower.rsplit(".", 1)[-1]


def _decode_plain(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        raise ValueError("Password-protected PDFs are not supported.")
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text()
        except Exception as e:  # noqa: BLE001
            logger.debug("PDF page extract failed: %s", e)
            continue
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)


def _extract_html(data: bytes) -> str:
    from app.services.website_ingest import extract_visible_text

    html = _decode_plain(data)
    return extract_visible_text(html)


def _extract_xml_text(data: bytes) -> str:
    raw = _decode_plain(data)
    try:
        root = ET.fromstring(data)
        text = " ".join(root.itertext())
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            return text
    except ET.ParseError:
        pass
    return re.sub(r"<[^>]+>", " ", raw)


def _extract_json(data: bytes) -> str:
    raw = _decode_plain(data)
    try:
        obj = json.loads(raw)
    except json.JSONDecodeError:
        return raw

    def walk(o: object) -> list[str]:
        out: list[str] = []
        if isinstance(o, dict):
            for k, v in o.items():
                ks = str(k).strip()
                if ks:
                    out.append(ks)
                out.extend(walk(v))
        elif isinstance(o, list):
            for item in o:
                out.extend(walk(item))
        elif isinstance(o, bool):
            out.append("true" if o else "false")
        elif isinstance(o, (str, int, float)):
            s = str(o).strip()
            if s:
                out.append(s)
        return out

    parts = walk(obj)
    return "\n".join(parts)


def _extract_rtf(data: bytes) -> str:
    from striprtf.striprtf import rtf_to_text

    return rtf_to_text(_decode_plain(data))


def _extract_xlsx(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    try:
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" ".join(cells))
    finally:
        wb.close()
    return "\n".join(lines)


def _extract_xls(data: bytes) -> str:
    import xlrd

    book = xlrd.open_workbook(file_contents=data)
    lines: list[str] = []
    for sheet in book.sheets():
        for r in range(sheet.nrows):
            row_vals = sheet.row_values(r)
            cells = [str(c).strip() for c in row_vals if c is not None and str(c).strip()]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(lines)


def _extract_iwork_pages(data: bytes) -> str:
    """Apple Pages (.pages / .page) is a zip; often includes QuickLook/Preview.pdf."""
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile:
        return ""

    preview_pdfs = [
        n
        for n in zf.namelist()
        if n.lower().endswith("preview.pdf") or n.lower().endswith("/quicklook/preview.pdf")
    ]
    for name in sorted(preview_pdfs, key=len):
        try:
            pdf_bytes = zf.read(name)
            text = _extract_pdf(pdf_bytes)
            if text.strip():
                return text
        except Exception as e:  # noqa: BLE001
            logger.debug("Pages preview PDF read failed %s: %s", name, e)

    # Older single-file XML bundle
    for candidate in ("index.xml", "Index.xml"):
        if candidate in zf.namelist():
            try:
                xml_text = _extract_xml_text(zf.read(candidate))
                if xml_text.strip():
                    return xml_text
            except Exception as e:  # noqa: BLE001
                logger.debug("Pages index.xml failed: %s", e)
    return ""


def extract_plain_text(filename: str, data: bytes) -> str:
    """
    Best-effort text extraction. Raises ValueError for unknown/unsupported types
    or formats we deliberately skip (.doc binary).
    """
    ext = normalize_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext or 'unknown'}' for {filename!r}")

    if ext in (".txt", ".text", ".md", ".markdown", ".csv", ".tsv"):
        return _decode_plain(data)

    if ext in (".html", ".htm"):
        return _extract_html(data)

    if ext == ".xml":
        return _extract_xml_text(data)

    if ext == ".json":
        return _extract_json(data)

    if ext == ".rtf":
        return _extract_rtf(data)

    if ext == ".pdf":
        return _extract_pdf(data)

    if ext == ".docx":
        return _extract_docx(data)

    if ext == ".doc":
        raise ValueError(
            "Legacy Word .doc is not supported. Save as .docx or export PDF and upload again."
        )

    if ext == ".xlsx":
        return _extract_xlsx(data)

    if ext == ".xls":
        return _extract_xls(data)

    if ext in (".pages", ".page"):
        text = _extract_iwork_pages(data)
        if not text.strip():
            raise ValueError(
                "Could not read this Pages document. Export as PDF or Word from Pages and upload that file."
            )
        return text

    raise ValueError(f"Unhandled extension {ext}")
